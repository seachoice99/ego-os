"""Tool Framework (v0.2, first capability).

The general mechanism by which an employee is granted a specific external
capability without ever holding a credential directly, per the Boundaries
section of architecture/005_EMPLOYEE_MODEL.md: an employee references a tool
by name; this module is the only place that resolves what that name actually
does and whether the employee's declared `permissions` allow it.

Adding a new tool later (web search, document generation, spreadsheet
editing) means adding an entry to TOOLS, not changing this framework.
"""

import html
import os
import re
import shutil
import zipfile
from pathlib import Path

import fitz  # PyMuPDF
import httpx
from docx import Document
from fpdf import FPDF
from openpyxl import Workbook
from openpyxl.styles import Font
from PIL import Image

REPO_ROOT = Path(__file__).parent.parent.resolve()
GENERATED_DIR = Path(__file__).parent / "generated"
UPLOADS_DIR = Path(__file__).parent / "uploads"
PRESENTATIONS_DIR = Path(
    os.environ.get("PRESENTATIONS_DIR", str(Path(__file__).parent / "generated" / "_presentations"))
)
TAVILY_SEARCH_URL = "https://api.tavily.com/search"
_DOCUMENT_FORMATS = {".md", ".docx", ".pdf"}

# Paths a tool must never touch, even for an employee with write_repository --
# these hold credentials or version-control internals, not task output.
_DENYLIST = {".env", ".git"}


class ToolError(Exception):
    """Raised when a tool cannot or must not run: bad arguments, a path
    outside the repository, or a denylisted path."""


def _resolve_repo_path(path: str) -> Path:
    target = (REPO_ROOT / path).resolve()
    if not target.is_relative_to(REPO_ROOT):
        raise ToolError(f"path '{path}' is outside the repository and cannot be accessed")
    if target.name in _DENYLIST or any(part in _DENYLIST for part in target.relative_to(REPO_ROOT).parts):
        raise ToolError(f"path '{path}' is not accessible to tools")
    return target


def _read_repository_file(path: str) -> str:
    target = _resolve_repo_path(path)
    if not target.is_file():
        raise ToolError(f"no such file: {path}")
    return target.read_text(encoding="utf-8")


def _write_repository_file(path: str, content: str) -> str:
    target = _resolve_repo_path(path)
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(content, encoding="utf-8")
    return f"wrote {len(content)} characters to {path}"


def _web_search(query: str) -> str:
    api_key = os.environ.get("TAVILY_API_KEY")
    if not api_key:
        raise ToolError("TAVILY_API_KEY is not configured in .env")

    response = httpx.post(
        TAVILY_SEARCH_URL,
        json={"api_key": api_key, "query": query, "max_results": 5},
        timeout=20,
    )
    response.raise_for_status()
    results = response.json().get("results", [])
    if not results:
        return "No results found."

    lines = []
    for r in results:
        title = r.get("title", "(no title)")
        url = r.get("url", "")
        snippet = r.get("content", "")
        lines.append(f"- {title} ({url})\n  {snippet}")
    return "\n".join(lines)


def _write_markdown(target: Path, content: str) -> None:
    target.write_text(content, encoding="utf-8")


def _write_docx(target: Path, content: str) -> None:
    """Minimal markdown-ish -> docx conversion: '# '/'## ' headings, '- '
    bullets, everything else a paragraph. Not a full markdown renderer --
    enough structure for a real, readable Word document."""
    doc = Document()
    for line in content.splitlines():
        stripped = line.strip()
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped:
            doc.add_paragraph(stripped)
    doc.save(target)


_UNICODE_TO_LATIN1 = {
    "—": "-", "–": "-",  # em/en dash
    "‘": "'", "’": "'",  # curly single quotes
    "“": '"', "”": '"',  # curly double quotes
    "…": "...",
    "•": "-",  # bullet
}


def _latin1_safe(text: str) -> str:
    """The PDF core font (Helvetica) only supports Latin-1/WinAnsi, but
    LLM output routinely contains em-dashes, curly quotes, and bullet
    characters -- found live during verification (raised
    FPDFUnicodeEncodingException). Transliterate the common ones, then
    drop anything still unsupported rather than crashing generation."""
    for unicode_char, replacement in _UNICODE_TO_LATIN1.items():
        text = text.replace(unicode_char, replacement)
    return text.encode("latin-1", errors="replace").decode("latin-1")


def _write_pdf(target: Path, content: str) -> None:
    """Same minimal heading/bullet/paragraph structure as _write_docx,
    rendered with the PDF core font (latin-1 only, so a plain '-' bullet
    rather than a unicode glyph).

    Every multi_cell call must explicitly return the cursor to the left
    margin (new_x="LMARGIN"): fpdf2 defaults new_x to the *right* edge, so
    without this the next call starts with almost no width left on the
    line and raises "Not enough horizontal space to render a single
    character" -- found live during Document Generation verification."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    for line in content.splitlines():
        stripped = _latin1_safe(line.strip())
        if stripped.startswith("## "):
            pdf.set_font("Helvetica", "B", 14)
            pdf.multi_cell(0, 8, stripped[3:], new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", size=12)
        elif stripped.startswith("# "):
            pdf.set_font("Helvetica", "B", 16)
            pdf.multi_cell(0, 10, stripped[2:], new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("Helvetica", size=12)
        elif stripped.startswith("- "):
            pdf.multi_cell(0, 8, f"- {stripped[2:]}", new_x="LMARGIN", new_y="NEXT")
        elif stripped:
            pdf.multi_cell(0, 8, stripped, new_x="LMARGIN", new_y="NEXT")
        else:
            pdf.ln(4)
    pdf.output(str(target))


def _create_document(filename: str, content: str, task_id: int) -> str:
    if "/" in filename or "\\" in filename or ".." in filename:
        raise ToolError("filename must not contain path separators")
    ext = Path(filename).suffix.lower()
    if ext not in _DOCUMENT_FORMATS:
        raise ToolError(f"unsupported document format '{ext or filename}' -- use .md, .docx, or .pdf")

    task_dir = GENERATED_DIR / str(task_id)
    task_dir.mkdir(parents=True, exist_ok=True)
    target = task_dir / filename

    if ext == ".md":
        _write_markdown(target, content)
    elif ext == ".docx":
        _write_docx(target, content)
    elif ext == ".pdf":
        _write_pdf(target, content)

    return f"created document artifact '{filename}', downloadable from the task's report."


def _create_spreadsheet(filename: str, data, task_id: int) -> str:
    if "/" in filename or "\\" in filename or ".." in filename:
        raise ToolError("filename must not contain path separators")
    if Path(filename).suffix.lower() != ".xlsx":
        raise ToolError(f"unsupported spreadsheet format '{filename}' -- use .xlsx")
    if not isinstance(data, list) or not data or not all(isinstance(row, list) for row in data):
        raise ToolError('data must be a non-empty list of rows, e.g. [["Header1","Header2"],["a","b"]]')

    task_dir = GENERATED_DIR / str(task_id)
    task_dir.mkdir(parents=True, exist_ok=True)
    target = task_dir / filename

    wb = Workbook()
    ws = wb.active
    for row in data:
        ws.append(row)
    for cell in ws[1]:
        cell.font = Font(bold=True)
    for column_cells in ws.columns:
        widest = max((len(str(cell.value)) for cell in column_cells if cell.value is not None), default=10)
        ws.column_dimensions[column_cells[0].column_letter].width = min(widest + 2, 60)
    wb.save(target)

    return f"created spreadsheet artifact '{filename}' with {len(data)} rows, downloadable from the task's report."


_SITE_SLUG_RE = re.compile(r"^[a-z0-9][a-z0-9-]{1,48}$")
_HEX_COLOR_RE = re.compile(r"^#[0-9a-fA-F]{3}$|^#[0-9a-fA-F]{6}$")
_SLIDE_IMAGE_EXTS = {".png", ".jpg", ".jpeg"}
_DEFAULT_ACCENT = "#3b82f6"

_PRESENTATION_HTML = """<!doctype html>
<html>
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{title}</title>
<style>:root {{ --accent: {accent}; }}</style>
<link rel="stylesheet" href="styles.css">
</head>
<body>
<div class="viewer-shell">
  <aside class="thumb-panel">
{thumbs}
  </aside>
  <div class="main-viewer">
    <div class="deck" id="deck">
{slides}
    </div>
    <div class="deck-nav">
      <button type="button" class="deck-nav-btn" id="deck-prev" aria-label="Previous slide">&#9650;</button>
      <span class="deck-counter"><span id="deck-current">1</span> / {count}</span>
      <button type="button" class="deck-nav-btn" id="deck-next" aria-label="Next slide">&#9660;</button>
    </div>
  </div>
</div>
<script src="script.js"></script>
</body>
</html>
"""

_PRESENTATION_CSS = """* { box-sizing: border-box; }
html, body { margin: 0; padding: 0; height: 100%; background: #0b0d12; color: #e8e8ec; font-family: system-ui, sans-serif; }
.viewer-shell { display: flex; height: 100vh; }
.thumb-panel {
  width: clamp(140px, 15vw, 220px); flex-shrink: 0; overflow-y: auto;
  background: #111319; border-right: 1px solid #22262f; padding: 0.8rem 0.6rem;
  display: flex; flex-direction: column; gap: 0.6rem; justify-content: space-evenly;
}
.thumb {
  display: block; width: 100%; aspect-ratio: 16 / 9; flex-shrink: 0;
  background: none; border: 2px solid transparent; border-radius: 6px;
  padding: 0; cursor: pointer; position: relative; overflow: hidden;
}
.thumb img { width: 100%; height: 100%; object-fit: cover; display: block; opacity: 0.55; transition: opacity 0.15s; }
.thumb span {
  position: absolute; top: 4px; left: 4px; font-size: 0.8rem; font-weight: 700;
  color: #fff; background: rgba(0,0,0,0.7); padding: 0.05rem 0.4rem; border-radius: 4px;
}
.thumb:hover img, .thumb.active img { opacity: 1; }
.thumb.active { border-color: var(--accent); }
.thumb.active span { background: var(--accent); }
.main-viewer { flex: 1; overflow-y: scroll; scroll-snap-type: y proximity; position: relative; }
.slide-frame { scroll-snap-align: start; min-height: 100vh; display: flex; align-items: center; justify-content: center; padding: 1.5rem; }
.slide { max-width: 100%; text-align: center; }
.slide img { max-width: 100%; max-height: 92vh; display: block; margin: 0 auto; border-radius: 4px; box-shadow: 0 10px 40px rgba(0,0,0,0.5); }
.slide-caption { color: #aab; margin-top: 0.8rem; font-size: 0.95rem; }
.deck-nav {
  position: fixed; right: 1rem; bottom: 1rem; display: flex; align-items: center; gap: 0.5rem;
  background: rgba(0,0,0,0.55); border-radius: 999px; padding: 0.3rem 0.5rem;
}
.deck-nav-btn {
  width: 1.8rem; height: 1.8rem; border-radius: 50%; border: none; cursor: pointer;
  background: rgba(255,255,255,0.08); color: var(--accent); font-size: 0.8rem;
  display: flex; align-items: center; justify-content: center; line-height: 1;
}
.deck-nav-btn:hover { background: var(--accent); color: #0b0d12; }
.deck-nav-btn:disabled { opacity: 0.3; cursor: default; background: rgba(255,255,255,0.08); color: var(--accent); }
.deck-counter { color: var(--accent); font-size: 0.85rem; font-weight: 600; padding: 0 0.2rem; }
::-webkit-scrollbar { width: 8px; }
::-webkit-scrollbar-thumb { background: var(--accent); border-radius: 4px; }
"""

_PRESENTATION_JS = """(function () {
  var counter = document.getElementById('deck-current');
  var prevBtn = document.getElementById('deck-prev');
  var nextBtn = document.getElementById('deck-next');
  var thumbs = Array.prototype.slice.call(document.querySelectorAll('.thumb'));
  var slides = Array.prototype.slice.call(document.querySelectorAll('.slide-frame'));
  var current = 1;

  function goTo(index) {
    index = Math.max(1, Math.min(slides.length, index));
    var target = slides[index - 1];
    if (target) target.scrollIntoView({ behavior: 'smooth', block: 'start' });
  }

  function updateActive(index) {
    current = index;
    counter.textContent = index;
    thumbs.forEach(function (t) { t.classList.remove('active'); });
    if (thumbs[index - 1]) {
      thumbs[index - 1].classList.add('active');
      thumbs[index - 1].scrollIntoView({ block: 'nearest' });
    }
    prevBtn.disabled = index <= 1;
    nextBtn.disabled = index >= slides.length;
  }

  thumbs.forEach(function (btn) {
    btn.addEventListener('click', function () {
      var target = document.getElementById(btn.getAttribute('data-target'));
      if (target) target.scrollIntoView({ behavior: 'smooth' });
    });
  });

  prevBtn.addEventListener('click', function () { goTo(current - 1); });
  nextBtn.addEventListener('click', function () { goTo(current + 1); });

  document.addEventListener('keydown', function (e) {
    if (e.key === 'ArrowUp' || e.key === 'PageUp') { e.preventDefault(); goTo(current - 1); }
    if (e.key === 'ArrowDown' || e.key === 'PageDown' || e.key === ' ') { e.preventDefault(); goTo(current + 1); }
  });

  if ('IntersectionObserver' in window) {
    var observer = new IntersectionObserver(function (entries) {
      entries.forEach(function (entry) {
        if (entry.isIntersecting) {
          updateActive(slides.indexOf(entry.target) + 1);
        }
      });
    }, { threshold: 0.6 });
    slides.forEach(function (s) { observer.observe(s); });
  }

  updateActive(1);
})();
"""


def _build_presentation_site(site_name: str, captions, task_id: int, accent: str = None) -> str:
    """Build and publish a real, browsable presentation website from this
    task's uploaded slide deck -- a .zip of images, or a .pdf (each page
    rendered to an image) -- per architecture/007's MVP slice: dark theme,
    vertical scroll, thumbnail nav, deck counter (video hotspots/portfolio
    grid/derived-PDF-export deliberately deferred; that's the site being
    exported back to PDF, unrelated to accepting a PDF as input here).

    Deliberately one deterministic tool call rather than a multi-step
    agent loop: the image resize / HTML generation / publish steps are
    ordinary code, not something that needs an LLM choreographing several
    tool calls per turn, so the existing one-tool-call-per-turn execution
    model (ego_os/lifecycle.py) did not need to change to support this."""
    if not _SITE_SLUG_RE.match(site_name or ""):
        raise ToolError("site_name must be lowercase letters, digits, and hyphens only (2-49 chars)")
    if not isinstance(captions, list) or not all(isinstance(c, str) for c in captions):
        raise ToolError(
            'captions must be a JSON array of strings, one per slide in file order '
            '(use "" for none), e.g. ["Cover", "Agenda", ""]'
        )
    accent = accent or _DEFAULT_ACCENT
    if not _HEX_COLOR_RE.match(accent):
        raise ToolError("accent must be a hex color like #3b82f6")

    task_uploads = UPLOADS_DIR / str(task_id)
    sources = sorted(
        p for p in task_uploads.iterdir() if p.suffix.lower() in (".zip", ".pdf")
    ) if task_uploads.is_dir() else []
    if not sources:
        raise ToolError(
            "no uploaded slide deck found for this task -- the Owner must attach a .zip of "
            "slide images (.png/.jpg/.jpeg) or a .pdf deck when submitting the task"
        )

    source_dir = GENERATED_DIR / str(task_id) / "source"
    source_dir.mkdir(parents=True, exist_ok=True)
    source_dir_resolved = source_dir.resolve()
    source_file = sources[0]

    if source_file.suffix.lower() == ".zip":
        with zipfile.ZipFile(source_file) as zf:
            for member in zf.namelist():
                if Path(member).suffix.lower() not in _SLIDE_IMAGE_EXTS:
                    continue
                member_path = (source_dir / Path(member).name).resolve()
                if not member_path.is_relative_to(source_dir_resolved):
                    continue  # zip-slip guard: never extract outside source_dir
                with zf.open(member) as src, open(member_path, "wb") as dst:
                    shutil.copyfileobj(src, dst)
    else:
        # .pdf: render each page to a PNG at 2x scale (~144 DPI) -- plenty
        # of source resolution for the later resize-to-1600px-wide step,
        # without the huge intermediate files a higher zoom would produce.
        with fitz.open(source_file) as pdf:
            zoom = fitz.Matrix(2, 2)
            for page_index, page in enumerate(pdf, start=1):
                pixmap = page.get_pixmap(matrix=zoom)
                pixmap.save(source_dir / f"pdfpage{page_index:03d}.png")

    images = sorted(p for p in source_dir.iterdir() if p.suffix.lower() in _SLIDE_IMAGE_EXTS)
    if not images:
        raise ToolError("uploaded slide deck contained no usable slides")

    site_dir = GENERATED_DIR / str(task_id) / "site"
    img_dir = site_dir / "img"
    img_dir.mkdir(parents=True, exist_ok=True)

    slide_parts, thumb_parts = [], []
    for index, image_path in enumerate(images, start=1):
        caption = captions[index - 1] if index - 1 < len(captions) else ""
        out_name = f"s{index:03d}.jpg"
        with Image.open(image_path) as im:
            im = im.convert("RGB")
            if im.width > 1600:
                new_height = round(im.height * (1600 / im.width))
                im = im.resize((1600, new_height), Image.LANCZOS)
            im.save(img_dir / out_name, "JPEG", quality=80, optimize=True)

        safe_caption = html.escape(caption)
        alt = safe_caption or f"Slide {index}"
        caption_html = f'<p class="slide-caption">{safe_caption}</p>' if safe_caption else ""
        loading = "eager" if index <= 2 else "lazy"
        slide_parts.append(
            f'      <section class="slide-frame" id="slide-{index}">\n'
            f'        <div class="slide"><img src="img/{out_name}" alt="{alt}" loading="{loading}">'
            f'{caption_html}</div>\n      </section>'
        )
        thumb_parts.append(
            f'    <button type="button" class="thumb" data-target="slide-{index}">'
            f'<img src="img/{out_name}" alt=""><span>{index}</span></button>'
        )

    (site_dir / "index.html").write_text(
        _PRESENTATION_HTML.format(
            title=html.escape(site_name),
            accent=accent,
            thumbs="\n".join(thumb_parts),
            slides="\n".join(slide_parts),
            count=len(images),
        ),
        encoding="utf-8",
    )
    (site_dir / "styles.css").write_text(_PRESENTATION_CSS, encoding="utf-8")
    (site_dir / "script.js").write_text(_PRESENTATION_JS, encoding="utf-8")

    publish_dir = PRESENTATIONS_DIR / site_name
    PRESENTATIONS_DIR.mkdir(parents=True, exist_ok=True)
    if publish_dir.exists():
        shutil.rmtree(publish_dir)
    shutil.copytree(site_dir, publish_dir)

    return (
        f"published presentation site '{site_name}' with {len(images)} slides, "
        f"reachable at /p/{site_name}/."
    )


TOOLS = {
    "read_repository_file": {
        "permission": "read_repository",
        "description": (
            'read_repository_file(path): read a text file from this repository. '
            'Args as JSON: {"path": "relative/path.ext"}'
        ),
        "fn": _read_repository_file,
    },
    "write_repository_file": {
        "permission": "write_repository",
        "description": (
            'write_repository_file(path, content): create or overwrite a text file in this '
            'repository. Args as JSON: {"path": "relative/path.ext", "content": "..."}'
        ),
        "fn": _write_repository_file,
    },
    "web_search": {
        "permission": "use_web",
        "description": (
            'web_search(query): search the live web and return up to 5 results, each with a '
            'title, url, and content snippet. Args as JSON: {"query": "search terms"}'
        ),
        "fn": _web_search,
    },
    "create_document": {
        "permission": "create_documents",
        "needs_context": ["task_id"],
        "produces_artifact": "document",
        "description": (
            'create_document(filename, content): generate a real, downloadable document artifact. '
            "filename must end in .md, .docx, or .pdf. content is plain text where lines starting "
            "with '# ' or '## ' become headings and '- ' becomes a bullet. "
            'Args as JSON: {"filename": "name.pdf", "content": "# Title\\n\\nBody text..."}'
        ),
        "fn": _create_document,
    },
    "create_spreadsheet": {
        "permission": "create_finance_reports",
        "needs_context": ["task_id"],
        "produces_artifact": "spreadsheet",
        "description": (
            'create_spreadsheet(filename, data): generate a real, downloadable .xlsx spreadsheet. '
            "filename must end in .xlsx. data is a JSON array of rows, each row a JSON array of "
            "cell values; the first row is treated as the header and rendered bold. "
            'Args as JSON: {"filename": "report.xlsx", "data": [["Employee","Cost"],["writer",0.0123]]}'
        ),
        "fn": _create_spreadsheet,
    },
    "build_presentation_site": {
        "permission": "build_presentation_sites",
        "needs_context": ["task_id"],
        "produces_artifact": "website",
        "description": (
            "build_presentation_site(site_name, captions, accent): build and publish a real, "
            "browsable presentation website from this task's uploaded slide deck -- a .zip of "
            ".png/.jpg/.jpeg images, or a .pdf (each page becomes one slide) -- attached when "
            "the task was submitted (fails if nothing was attached). site_name is a URL slug "
            "(lowercase letters/digits/hyphens only). "
            'captions is a JSON array of strings, one per slide in file order (use "" for none). '
            "accent is an optional hex color for links/highlights (default #3b82f6). "
            'Args as JSON: {"site_name": "t1-tender-pitch", "captions": ["Cover", "Agenda", ""], '
            '"accent": "#e11d48"}'
        ),
        "fn": _build_presentation_site,
    },
}


def available_tools(permissions):
    """Tools an employee with this permission list is allowed to see and use."""
    return [tool for tool in TOOLS.values() if tool["permission"] in permissions]


def call_tool(permissions, tool_name: str, context=None, **kwargs):
    """context carries values a tool needs but an LLM's TOOL_REQUEST can't
    reasonably supply itself (e.g. which task a generated document belongs
    to) -- only merged in for a tool that declares it wants that key via
    `needs_context`, so existing tools are unaffected."""
    tool = TOOLS.get(tool_name)
    if tool is None:
        raise ToolError(f"unknown tool: {tool_name}")
    if tool["permission"] not in permissions:
        raise ToolError(
            f"tool '{tool_name}' requires permission '{tool['permission']}', which this employee does not have"
        )
    for key in tool.get("needs_context", []):
        if context and key in context:
            kwargs[key] = context[key]
    return tool["fn"](**kwargs)
